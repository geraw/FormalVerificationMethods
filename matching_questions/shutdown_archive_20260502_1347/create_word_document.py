"""
Create a RTL Word document for a matching question.

The six diagrams are rendered from a single TikZ source arranged as a 2x3
matrix: program graphs on the first row and transition systems on the second.
"""

from pathlib import Path
import subprocess

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(r"c:\Users\geraw\courses\FormalVerificationMethods\matching_questions")
DOCX_NAME = "MOED-ALEPH.docx"
TIKZ_TEX_NAME = "matching_diagrams_matrix.tex"
TIKZ_PNG_NAME = "matching_diagrams_matrix.png"


def _append_unique(parent, tag):
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    el = OxmlElement(tag)
    parent.append(el)
    return el


def _set_doc_rtl(doc):
    settings = doc.settings.element
    _append_unique(settings, "w:bidi")


def _set_table_rtl(table):
    tbl_pr = table._tbl.tblPr
    bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        tbl_pr.append(OxmlElement("w:bidiVisual"))


def _set_run_rtl(run):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    r_pr = run._element.get_or_add_rPr()
    if r_pr.find(qn("w:rtl")) is None:
        r_pr.append(OxmlElement("w:rtl"))


def set_hebrew_formatting(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p_pr = paragraph._element.get_or_add_pPr()
    _append_unique(p_pr, "w:bidi")
    jc = _append_unique(p_pr, "w:jc")
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        jc.set(qn("w:val"), "center")
    elif align == WD_ALIGN_PARAGRAPH.LEFT:
        jc.set(qn("w:val"), "left")
    else:
        jc.set(qn("w:val"), "right")
    for run in paragraph.runs:
        _set_run_rtl(run)


def _set_cell_border(cell, border_setting):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        cfg = border_setting.get("all") or border_setting.get(side)
        el = OxmlElement(f"w:{side}")
        if cfg:
            el.set(qn("w:val"), cfg.get("val", "single"))
            el.set(qn("w:sz"), str(cfg.get("sz", 6)))
            el.set(qn("w:color"), cfg.get("color", "000000"))
        else:
            el.set(qn("w:val"), "none")
        borders.append(el)
    tc_pr.append(borders)


def _add_run(paragraph, text, bold=False, italic=False, size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    _set_run_rtl(run)
    return run


def _clear_paragraph(paragraph):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def _add_rtl_paragraph(container, text="", bold=False, italic=False, size=12,
                       align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph = container.add_paragraph()
    if text:
        _add_run(paragraph, text, bold=bold, italic=italic, size=size)
    set_hebrew_formatting(paragraph, align)
    return paragraph


def _force_document_right_alignment(doc):
    """Apply right alignment and RTL to every text paragraph in the document."""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_hebrew_formatting(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
    for table in doc.tables:
        _set_table_rtl(table)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        set_hebrew_formatting(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)


def _write_tikz_source(output_dir):
    tex_path = output_dir / TIKZ_TEX_NAME
    tex = r"""\documentclass[tikz,border=6pt]{standalone}
\usetikzlibrary{arrows.meta,positioning,calc,fit,matrix}
\makeatletter
\AtBeginDocument{\global\let\pgf@selectfontorig\relax}
\makeatother

\tikzset{
  nodebox/.style={draw=#1, rounded corners=4pt, very thick, fill=gray!4,
    minimum width=12mm, minimum height=8mm, align=center, inner sep=2pt},
  nodebox/.default=black,
  initial/.style={-Latex, very thick},
  edge/.style={-Latex, very thick, shorten >=3pt, shorten <=3pt},
  labelbox/.style={fill=white, inner sep=1.2pt},
  aptxt/.style={text=gray!65!black, anchor=north west}
}

\newcommand{\initarrow}[2]{\draw[initial] #1 -- (#2);}
\begin{document}
\begin{tikzpicture}[scale=.82, transform shape]
  \begin{scope}[shift={(0,0)}]
    \node at (1.6,2.25) {\bfseries\large TS1};
    \node[nodebox] (q0) at (0,0) {$q_0$};
    \node[nodebox] (q1) at (3,0) {$q_1$};
    \node[nodebox] (q2) at (1.5,1.65) {$q_2$};
    \initarrow{(-.85,0)}{q0.west}
    \draw[edge] (q0) -- node[labelbox,above] {\scriptsize $x$} (q1);
    \draw[edge] (q1) -- node[labelbox,right] {\scriptsize $y$} (q2);
    \draw[edge] (q2) -- node[labelbox,left] {\scriptsize $\gamma$} (q0);
    \draw[edge,bend right=20] (q1) to node[labelbox,below] {\scriptsize $x$} (q0);
  \end{scope}

  \begin{scope}[shift={(5.3,0)}]
    \node at (1.55,2.25) {\bfseries\large TS2};
    \node[nodebox] (t0) at (0,0) {$t_0$};
    \node[nodebox] (t1) at (3.1,0) {$t_1$};
    \initarrow{(-.85,0)}{t0.west}
    \draw[edge,bend left=22] (t0) to node[labelbox,above] {\scriptsize $d$} (t1);
    \draw[edge,bend left=22] (t1) to node[labelbox,below] {\scriptsize $c$} (t0);
    \draw[edge,loop above,looseness=8] (t0) to node[labelbox,above] {\scriptsize $c$} (t0);
  \end{scope}

  \begin{scope}[shift={(10.5,0)}]
    \node at (1.55,2.25) {\bfseries\large TS3};
    \node[nodebox] (s0) at (0,0) {$s_0$};
    \node[nodebox] (s1) at (3.1,0) {$s_1$};
    \initarrow{(0,1.05)}{s0.north}
    \draw[edge,bend left=22] (s0) to node[labelbox,above] {\scriptsize $a$} (s1);
    \draw[edge,bend left=22] (s1) to node[labelbox,below] {\scriptsize $b$} (s0);
  \end{scope}

  \begin{scope}[shift={(0,4.6)}]
    \node at (1.6,2.25) {\bfseries\large PG1};
    \node[nodebox] (l0) at (0,0) {$\ell_0$};
    \node[nodebox] (l1) at (3.1,0) {$\ell_1$};
    \initarrow{(-.85,0)}{l0.west}
    \draw[edge,bend left=22] (l0) to node[labelbox,above] {\scriptsize $z:=1$} (l1);
    \draw[edge,bend left=22] (l1) to node[labelbox,below] {\scriptsize $z:=0$} (l0);
    \draw[edge,loop above,looseness=8] (l0) to node[labelbox,above] {\scriptsize $z:=0$} (l0);
  \end{scope}

  \begin{scope}[shift={(5.3,4.6)}]
    \node at (1.6,2.25) {\bfseries\large PG2};
    \node[nodebox] (l0) at (0,0) {$\ell_0$};
    \node[nodebox] (l1) at (3.2,0) {$\ell_1$};
    \initarrow{(0,1.05)}{l0.north}
    \draw[edge,bend left=22] (l0) to node[labelbox,above] {\scriptsize $x:=1$} (l1);
    \draw[edge,bend left=22] (l1) to node[labelbox,below] {\scriptsize $x:=x+1$} (l0);
  \end{scope}

  \begin{scope}[shift={(10.5,4.6)}]
    \node at (1.6,2.25) {\bfseries\large PG3};
    \node[nodebox] (l0) at (0,0) {$\ell_0$};
    \node[nodebox] (l1) at (3,1.45) {$\ell_1$};
    \node[nodebox] (l2) at (3,-1.45) {$\ell_2$};
    \initarrow{(-.85,0)}{l0.west}
    \draw[edge] (l0) -- node[labelbox,above left] {\scriptsize $y>0:\mathrm{skip}$} (l1);
    \draw[edge] (l1) -- node[labelbox,right] {\scriptsize $y:=y-1$} (l2);
    \draw[edge] (l2) -- node[labelbox,below left] {\scriptsize $y:=0$} (l0);
    \draw[edge,loop above,looseness=8] (l0) to node[labelbox,above] {\scriptsize $y\le0:\mathrm{skip}$} (l0);
  \end{scope}
\end{tikzpicture}
\end{document}
"""
    tex_path.write_text(tex, encoding="utf-8")
    return tex_path


def _render_tikz_matrix(output_dir):
    tex_path = output_dir / TIKZ_TEX_NAME
    if not tex_path.exists():
        raise FileNotFoundError(f"TikZ source not found: {tex_path}")
    subprocess.run(
        ["pdflatex", "-interaction=nonstopmode", tex_path.name],
        cwd=output_dir,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    subprocess.run(
        ["pdftoppm", "-png", "-r", "300", TIKZ_TEX_NAME.replace(".tex", ".pdf"), "matching_diagrams_matrix"],
        cwd=output_dir,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    generated = output_dir / "matching_diagrams_matrix-1.png"
    final = output_dir / TIKZ_PNG_NAME
    if final.exists():
        final.unlink()
    generated.rename(final)
    return final


def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


def _configure_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:cs"), "Arial")
    p_pr = style.element.get_or_add_pPr()
    _append_unique(p_pr, "w:bidi")
    jc = _append_unique(p_pr, "w:jc")
    jc.set(qn("w:val"), "right")


def create_tomax_word_document(output_dir="."):
    output_dir = Path(output_dir)
    matrix_png = _render_tikz_matrix(output_dir)

    doc = Document()
    _set_doc_rtl(doc)
    _configure_normal_style(doc)

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_rtl(hdr_table)
    _remove_table_borders(hdr_table)

    title_cell = hdr_table.cell(0, 0)
    info_cell = hdr_table.cell(0, 1)

    _clear_paragraph(title_cell.paragraphs[0])
    _add_run(title_cell.paragraphs[0], "בוחן אמצע בקורס", bold=True, size=20)
    set_hebrew_formatting(title_cell.paragraphs[0])
    subtitle = title_cell.add_paragraph()
    _add_run(subtitle, "'מבוא לאימות תוכנה בשיטות פורמליות'", bold=True, size=16)
    set_hebrew_formatting(subtitle)

    _set_cell_border(info_cell, {
        "top": {"val": "single", "sz": 6, "color": "000000"},
        "left": {"val": "single", "sz": 6, "color": "000000"},
        "bottom": {"val": "single", "sz": 6, "color": "000000"},
        "right": {"val": "single", "sz": 6, "color": "000000"},
    })
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    info_rows = [
        ("תאריך הבחינה:", "8.5.25"),
        ("שם המרצה:", "גרא וייס"),
        ("מספר הקורס:", "202-1-3061"),
        ("משך הבחינה:", "שלוש שעות"),
        ("חומר עזר:", "כל חומר עזר"),
    ]
    for index, (label, value) in enumerate(info_rows):
        paragraph = info_cell.paragraphs[0] if index == 0 else info_cell.add_paragraph()
        _clear_paragraph(paragraph)
        _add_run(paragraph, f"{label} ", bold=True, size=10)
        _add_run(paragraph, value, size=10)
        set_hebrew_formatting(paragraph)
        paragraph.paragraph_format.space_after = Pt(2)

    instructions = _add_rtl_paragraph(
        doc,
        "לכל שאלה תשובה אחת נכונה. "
        "במקרה של התלבטות, בחרו בתשובה שנראית לכם יותר נכונה.",
        size=11,
    )
    instructions.paragraph_format.space_before = Pt(10)
    instructions.paragraph_format.space_after = Pt(4)

    question = _add_rtl_paragraph(
        doc,
        "1. התאימו בין מערכות המעברים לבין גרפי התוכנית.",
        bold=True,
        size=12,
    )
    question.paragraph_format.space_before = Pt(6)
    question.paragraph_format.space_after = Pt(6)

    note = _add_rtl_paragraph(
        doc,
        "הערה: מערכות המעברים מציגות רק את המצבים הנגישים, ושמות המצבים שונו כפי שלמדנו בשיעור "
        "(על פי הערכים של המשתנים בכל מצב). בנוסף, השמטנו את התוויות של המצבים במערכות המעברים.",
        italic=True,
        size=10,
    )
    note.paragraph_format.space_after = Pt(6)

    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_para.add_run()
    run.add_picture(str(matrix_png), width=Inches(6.35))
    image_para.paragraph_format.space_after = Pt(8)

    matches = [
        ("א", "TS1↔PG3, TS2↔PG1, TS3↔PG2"),
        ("ב", "TS1↔PG1, TS2↔PG2, TS3↔PG3"),
        ("ג", "TS1↔PG2, TS2↔PG3, TS3↔PG1"),
        ("ד", "TS1↔PG1, TS2↔PG3, TS3↔PG2"),
        ("ה", "TS1↔PG2, TS2↔PG1, TS3↔PG3"),
        ("ו", "TS1↔PG3, TS2↔PG2, TS3↔PG1"),
    ]
    for letter, answer_text in matches:
        opt = _add_rtl_paragraph(doc, f"{letter}. {answer_text}", size=11)
        opt.paragraph_format.space_after = Pt(2)

    _force_document_right_alignment(doc)

    output_file = output_dir / DOCX_NAME
    doc.save(output_file)
    print(f"Word document created: {output_file}")
    return str(output_file)


if __name__ == "__main__":
    create_tomax_word_document(ROOT)
