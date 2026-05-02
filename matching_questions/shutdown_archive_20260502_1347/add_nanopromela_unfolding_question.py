from datetime import datetime
from pathlib import Path
import shutil
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "MOED-ALEPH-updated.docx"
BACKUP_DIR = ROOT / "word_backups"
TEX_NAME = "nanopromela_unfolding_graphs.tex"
PNG_NAME = "nanopromela_unfolding_graphs.png"


PROGRAM = """do
:: x < 2 ->
     x := x + 1;
     if
     :: y = 0 -> y := 1
     :: y = 1 -> y := 0; z := z + 1
     fi
:: x = 2 ->
     atomic{ y := 1; x := x + 1 }
od;
z := z + 1"""


def _set_run_font(run, name="Arial", size=12, rtl=True):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    r_pr = run._element.get_or_add_rPr()
    rtl_el = r_pr.find(qn("w:rtl"))
    if rtl and rtl_el is None:
        r_pr.append(OxmlElement("w:rtl"))
    if not rtl and rtl_el is not None:
        r_pr.remove(rtl_el)


def _set_paragraph_direction(paragraph, rtl=True, align=None):
    p_pr = paragraph._element.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        p_pr.append(OxmlElement("w:bidi"))
    if not rtl and bidi is not None:
        p_pr.remove(bidi)
    if align is not None:
        paragraph.alignment = align


def add_rtl_paragraph(doc, text="", bold=False, size=12):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    _set_paragraph_direction(paragraph, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        _set_run_font(run, "Arial", size=size, rtl=True)
    return paragraph


def add_ltr_code(doc, text, size=9):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    _set_paragraph_direction(paragraph, rtl=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = paragraph.add_run(text)
    _set_run_font(run, "Courier New", size=size, rtl=False)
    return paragraph


def render_latex_graphs():
    subprocess.run(
        ["pdflatex", "-interaction=nonstopmode", TEX_NAME],
        cwd=ROOT,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    subprocess.run(
        ["pdftoppm", "-png", "-r", "300", TEX_NAME.replace(".tex", ".pdf"), "nanopromela_unfolding_graphs"],
        cwd=ROOT,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    generated = ROOT / "nanopromela_unfolding_graphs-1.png"
    final = ROOT / PNG_NAME
    if final.exists():
        final.unlink()
    generated.rename(final)
    return final


def main(target_docx=None):
    target = Path(target_docx) if target_docx else DOCX
    if not target.is_absolute():
        target = ROOT / target
    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = BACKUP_DIR / f"{target.stem}.before-nanopromela-graphs-{stamp}.docx"
    shutil.copy2(target, backup)

    matrix_png = render_latex_graphs()

    doc = Document(target)
    doc.add_page_break()
    add_rtl_paragraph(
        doc,
        "4. נתון קטע הקוד הבא ב-NanoPromela. איזה מבין גרפי התוכנית הבאים הוא הפריסה הנכונה של הקוד?",
        bold=True,
        size=12,
    )
    add_rtl_paragraph(
        doc,
        "הניחו כי x,y,z הם משתנים שלמים. שימו לב: בבחירת guarded command, בדיקת התנאי והפעולה האטומית הראשונה של הענף שנבחר מתבצעות כצעד אטומי אחד.",
        size=11,
    )
    add_ltr_code(doc, PROGRAM, size=9)
    add_rtl_paragraph(
        doc,
        "האפשרויות מסומנות באיור בשמות PG_1,...,PG_6.",
        size=11,
    )
    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(matrix_png), width=Inches(6.4))

    doc.save(target)
    print(f"updated={target}")
    print(f"backup={backup}")
    print(f"figure={matrix_png}")


if __name__ == "__main__":
    import sys

    main(sys.argv[1] if len(sys.argv) > 1 else None)
