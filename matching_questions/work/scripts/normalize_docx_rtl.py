from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "source.docx"

raise SystemExit(
    "Disabled: source.docx text and formatting are now edited manually in Word. "
    "Do not run this script unless you intentionally want to normalize all text."
)


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def mark_rtl_paragraph(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ensure_child(ppr, "w:bidi")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def mark_rtl_run(run) -> None:
    rpr = run._element.get_or_add_rPr()
    ensure_child(rpr, "w:rtl")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")


def paragraph_has_image(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//wp:inline//a:blip"))


def main() -> None:
    doc = Document(DOCX)
    changed = 0
    for paragraph in doc.paragraphs:
        if paragraph_has_image(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        mark_rtl_paragraph(paragraph)
        for run in paragraph.runs:
            mark_rtl_run(run)
        changed += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    mark_rtl_paragraph(paragraph)
                    for run in paragraph.runs:
                        mark_rtl_run(run)

    doc.save(DOCX)
    print(f"Normalized RTL/right alignment in {DOCX} ({changed} paragraphs)")


if __name__ == "__main__":
    main()
