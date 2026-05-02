from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

doc = Document(r"C:\Users\geraw\courses\FormalVerificationMethods\matching_questions\work\source.docx")

print("=== Q6 options 77-82 ===")
for i in range(77, 83):
    p = doc.paragraphs[i]
    text = p.text.strip()
    runs_detail = [(r.text, r.font.size, r.font.name) for r in p.runs]
    ppr = p._p.find(qn("w:pPr"))
    pb = ppr.find(qn("w:pageBreakBefore")) if ppr is not None else None
    print(f"[{i}] {repr(text[:40])}")
    print(f"     runs={runs_detail}")

print()
print("=== Question starts ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for q in ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9."]:
        if text.startswith(q + " ") or text.startswith(q + "\t"):
            ppr = p._p.find(qn("w:pPr"))
            pb_elem = ppr.find(qn("w:pageBreakBefore")) if ppr is not None else None
            has_pb = pb_elem is not None
            print(f"[{i}] pageBreakBefore={has_pb} {repr(text[:50])}")
