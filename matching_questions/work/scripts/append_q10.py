from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "source.docx"

raise SystemExit(
    "Disabled: source.docx text is now edited manually in Word. "
    "Do not run this script unless you intentionally want to recreate Q10 text."
)


def mark_rtl(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))


def set_para(paragraph, text: str, rtl: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    if rtl:
        mark_rtl(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    if rtl and rpr.find(qn("w:rtl")) is None:
        rpr.append(OxmlElement("w:rtl"))
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")


def clear_from_q10(doc: Document) -> None:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("10."):
            body = doc._body._element
            for para in list(doc.paragraphs[index:]):
                body.remove(para._p)
            return


def main() -> None:
    doc = Document(DOCX)
    clear_from_q10(doc)

    items = [
        "10. בשאלה זו נשתמש בכללי הגזירה הבאים להרכבת שתי מערכות מעברים.",
        "כללי הגזירה של שאלה 10:",
        "",
        "נתונות שתי מערכות המעברים של שאלה 10:",
        "",
        "סמנו את מערכת המעברים המתקבלת מן ההרכבה לפי כללי הגזירה של שאלה 10.",
        "אפשרויות הבחירה:",
        "א. אפשרות א",
        "",
        "ב. אפשרות ב",
        "",
        "ג. אפשרות ג",
        "",
        "ד. אפשרות ד",
        "",
        "ה. אפשרות ה",
        "",
        "ו. אפשרות ו",
        "",
    ]
    for text in items:
        paragraph = doc.add_paragraph()
        if text:
            set_para(paragraph, text)

    doc.save(DOCX)
    print(f"Appended Q10 to {DOCX}")


if __name__ == "__main__":
    main()
