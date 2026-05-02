from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "source.docx"

LETTERS = ["א", "ב", "ג", "ד", "ה", "ו"]

raise SystemExit(
    "Disabled: source.docx text is now edited manually in Word. "
    "Do not run this script unless you intentionally want to rewrite option labels."
)


def mark_rtl(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_text(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mark_rtl(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    if rpr.find(qn("w:rtl")) is None:
        rpr.append(OxmlElement("w:rtl"))
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")


def section_bounds(doc: Document, question: str, next_question: str) -> tuple[int, int]:
    start = None
    end = len(doc.paragraphs)
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith(question):
            start = index
        elif start is not None and text.startswith(next_question):
            end = index
            break
    if start is None:
        raise RuntimeError(f"Could not find section {question}")
    return start, end


def remove_inline_images(doc: Document, start: int, end: int, keep_first: int = 0) -> None:
    seen = 0
    for paragraph in doc.paragraphs[start:end]:
        if paragraph._p.xpath(".//wp:inline//a:blip"):
            seen += 1
            if seen > keep_first:
                clear_paragraph(paragraph)


def relabel_answers(doc: Document, start: int, end: int) -> None:
    option_index = None
    for index in range(start, end):
        if doc.paragraphs[index].text.strip() == "אפשרויות הבחירה:":
            option_index = index
            break
    if option_index is None:
        raise RuntimeError("Could not find answer header")

    label_index = 0
    for paragraph in doc.paragraphs[option_index + 1 : end]:
        text = paragraph.text.strip()
        if not text:
            continue
        if label_index >= len(LETTERS):
            break
        if text.startswith(f"{LETTERS[label_index]}."):
            letter = LETTERS[label_index]
            set_text(paragraph, f"{letter}. אפשרות {letter}")
            label_index += 1
    if label_index != len(LETTERS):
        raise RuntimeError(f"Expected 6 answer labels, found {label_index}")


def main() -> None:
    doc = Document(DOCX)

    q4_start, q4_end = section_bounds(doc, "4.", "5.")
    remove_inline_images(doc, q4_start, q4_end, keep_first=0)
    relabel_answers(doc, q4_start, q4_end)

    q6_start, q6_end = section_bounds(doc, "6.", "7.")
    remove_inline_images(doc, q6_start, q6_end, keep_first=1)
    relabel_answers(doc, q6_start, q6_end)

    doc.save(DOCX)
    print(f"Prepared {DOCX}")


if __name__ == "__main__":
    main()
