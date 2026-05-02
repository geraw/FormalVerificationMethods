from __future__ import annotations

import copy
from datetime import datetime
from pathlib import Path
import shutil
import subprocess
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "source.docx"
TEX_DIR = ROOT / "tex"
IMAGE_DIR = ROOT / "images"
BUILD_DIR = ROOT / "build"
BACKUP_DIR = ROOT / "backups"

PREVIEW_DPI = 600
EMBED_SVG_IN_DOCX = False
FULL_WIDTH = Inches(6.77)

QUESTION_IMAGES = [
    ("1.", "q01_transition_systems_vs_program_graphs", FULL_WIDTH),
    ("2.", "q02_interleaving_with_handshake", FULL_WIDTH),
    ("3.", "q03_logic_circuits_transition_systems", FULL_WIDTH),
    ("\u05d0. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d0", "q04_option_a", Inches(4.85)),
    ("\u05d1. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d1", "q04_option_b", Inches(4.85)),
    ("\u05d2. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d2", "q04_option_c", Inches(4.85)),
    ("\u05d3. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d3", "q04_option_d", Inches(4.85)),
    ("\u05d4. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d4", "q04_option_e", Inches(4.85)),
    ("\u05d5. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d5", "q04_option_f", Inches(4.85)),
    ("5.", "q05_async_channel_system", FULL_WIDTH),
    ("6.", "q06_alt_nanopromela_rules", FULL_WIDTH),
    ("\u05d0. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d0", "q06_option_a", Inches(4.85)),
    ("\u05d1. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d1", "q06_option_b", Inches(4.85)),
    ("\u05d2. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d2", "q06_option_c", Inches(4.85)),
    ("\u05d3. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d3", "q06_option_d", Inches(4.85)),
    ("\u05d4. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d4", "q06_option_e", Inches(4.85)),
    ("\u05d5. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d5", "q06_option_f", Inches(4.85)),
    ("7.", "q07_combined", FULL_WIDTH),
    ("8.", "q08_combined", FULL_WIDTH),
    ("\u05db\u05dc\u05dc\u05d9 \u05d4\u05d2\u05d6\u05d9\u05e8\u05d4 \u05d4\u05d7\u05d3\u05e9\u05d9\u05dd:", "q09_rules", FULL_WIDTH),
    ("\u05e0\u05ea\u05d5\u05e0\u05d5\u05ea \u05e9\u05ea\u05d9 \u05de\u05e2\u05e8\u05db\u05d5\u05ea \u05d4\u05de\u05e2\u05d1\u05e8\u05d9\u05dd", "q09_ts_input", FULL_WIDTH),
    ("\u05d0. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d0", "q09_option_a", Inches(4.35)),
    ("\u05d1. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d1", "q09_option_b", Inches(4.35)),
    ("\u05d2. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d2", "q09_option_c", Inches(4.35)),
    ("\u05d3. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d3", "q09_option_d", Inches(4.35)),
    ("\u05d4. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d4", "q09_option_e", Inches(4.35)),
    ("\u05d5. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d5", "q09_option_f", Inches(4.35)),
    ("10.", "q10_rules", FULL_WIDTH),
    ("\u05e0\u05ea\u05d5\u05e0\u05d5\u05ea \u05e9\u05ea\u05d9 \u05de\u05e2\u05e8\u05db\u05d5\u05ea \u05d4\u05de\u05e2\u05d1\u05e8\u05d9\u05dd:", "q10_ts_input", FULL_WIDTH),
    ("\u05d0. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d0", "q10_option_a", Inches(4.35)),
    ("\u05d1. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d1", "q10_option_b", Inches(4.35)),
    ("\u05d2. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d2", "q10_option_c", Inches(4.35)),
    ("\u05d3. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d3", "q10_option_d", Inches(4.35)),
    ("\u05d4. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d4", "q10_option_e", Inches(4.35)),
    ("\u05d5. \u05d0\u05e4\u05e9\u05e8\u05d5\u05ea \u05d5", "q10_option_f", Inches(4.35)),
]

SHARED_MEDIA_IMAGES = {}

ANSWER_LETTERS = ("א", "ב", "ג", "ד")


def run_command(args: list[str], cwd: Path) -> None:
    completed = subprocess.run(
        args,
        cwd=cwd,
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if completed.returncode != 0:
        command = " ".join(args)
        raise RuntimeError(f"Command failed: {command}\n{completed.stdout}")


def render_tex(stem: str) -> tuple[Path, Path]:
    tex_path = TEX_DIR / f"{stem}.tex"
    if not tex_path.exists():
        raise FileNotFoundError(f"Missing TeX source: {tex_path}")

    BUILD_DIR.mkdir(exist_ok=True)
    run_command(
        ["pdflatex", "-interaction=nonstopmode", f"-output-directory={BUILD_DIR}", tex_path.name],
        cwd=TEX_DIR,
    )

    pdf_path = BUILD_DIR / f"{stem}.pdf"
    if not pdf_path.exists():
        raise FileNotFoundError(f"pdflatex did not create {pdf_path}")

    IMAGE_DIR.mkdir(exist_ok=True)

    svg_path = IMAGE_DIR / f"{stem}.svg"
    run_command(["pdftocairo", "-svg", str(pdf_path), str(svg_path)], cwd=ROOT)
    if not svg_path.exists():
        raise FileNotFoundError(f"pdftocairo did not create {svg_path}")

    # PNG is used for the DOCX payload because some Word/Tomax import paths
    # drop SVG images entirely. TeX/PDF/SVG are still generated as the source
    # artifacts, while the Word document gets a high-resolution raster image.
    output_base = IMAGE_DIR / stem
    run_command(["pdftoppm", "-png", "-r", str(PREVIEW_DPI), str(pdf_path), str(output_base)], cwd=ROOT)

    generated = IMAGE_DIR / f"{stem}-1.png"
    final = IMAGE_DIR / f"{stem}.png"
    if final.exists():
        final.unlink()
    generated.rename(final)
    return svg_path, final


def inline_image_paragraphs(doc: Document):
    slots = []
    for paragraph in doc.paragraphs:
        if paragraph._p.xpath(".//wp:inline//a:blip"):
            slots.append(paragraph)
    return slots


def question_start_index(doc: Document, question_number: int) -> int:
    prefix = f"{question_number}."
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return index
    return 0


def find_inline_slot_after(doc: Document, anchor: str, used: set[int], start_index: int = 0):
    anchor_index = None
    fallback_prefix = None
    if "אפשרות" in anchor and len(anchor) >= 2:
        fallback_prefix = anchor[0] + "."

    for index, paragraph in enumerate(doc.paragraphs[start_index:], start=start_index):
        text = paragraph.text.strip()
        if anchor in paragraph.text or (fallback_prefix and text.startswith(fallback_prefix)):
            anchor_index = index
            break

    if anchor_index is None:
        raise RuntimeError(f"Could not find image anchor text: {anchor}")

    for index in range(anchor_index + 1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[index]
        if paragraph.text.strip():
            break
        if id(paragraph) in used:
            continue
        if paragraph._p.xpath(".//wp:inline//a:blip"):
            used.add(id(paragraph))
            return index, paragraph

    # No existing image slot found — insert a new empty paragraph right after the anchor
    anchor_para = doc.paragraphs[anchor_index]
    new_p = OxmlElement("w:p")
    anchor_para._p.addnext(new_p)
    for paragraph in doc.paragraphs:
        if paragraph._p is new_p:
            used.add(id(paragraph))
            return anchor_index + 1, paragraph
    raise RuntimeError(f"Inserted paragraph not found after anchor: {anchor}")


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def all_text(paragraph) -> str:
    return "".join(paragraph._p.xpath('.//*[local-name()="t"]/text()')).strip()


def mark_rtl(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))


def prepend_answer_label(paragraph, label: str) -> None:
    if paragraph.text.strip().startswith(f"{label}."):
        return

    run = paragraph.add_run(f"{label}. ")
    rpr = run._element.get_or_add_rPr()
    if rpr.find(qn("w:rtl")) is None:
        rpr.append(OxmlElement("w:rtl"))
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")

    paragraph._p.remove(run._r)
    paragraph._p.insert(0, run._r)
    mark_rtl(paragraph)


def ensure_formula_answer_labels(doc: Document) -> None:
    prompt_indexes = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if "סמנו את התשובה הנכונה" in all_text(paragraph)
    ]

    for prompt_index in prompt_indexes:
        candidates = []
        for paragraph in doc.paragraphs[prompt_index + 1 : prompt_index + 12]:
            text = all_text(paragraph)
            if not text:
                continue

            visible_text = paragraph.text.strip()
            already_labeled = any(visible_text.startswith(f"{letter}.") for letter in ANSWER_LETTERS)
            if visible_text and not already_labeled:
                continue

            candidates.append(paragraph)
            if len(candidates) == len(ANSWER_LETTERS):
                break

        if len(candidates) != len(ANSWER_LETTERS):
            raise RuntimeError(
                f"Expected {len(ANSWER_LETTERS)} formula answer paragraphs after prompt "
                f"at paragraph {prompt_index}, found {len(candidates)}"
            )

        for label, paragraph in zip(ANSWER_LETTERS, candidates):
            prepend_answer_label(paragraph, label)


def ensure_q6_distractor_answers(doc: Document) -> None:
    q6_index = None
    q7_index = len(doc.paragraphs)
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("6."):
            q6_index = index
        elif q6_index is not None and text.startswith("7."):
            q7_index = index
            break

    if q6_index is None:
        raise RuntimeError("Could not find Q6")

    answer_index = None
    for index in range(q6_index, q7_index):
        if doc.paragraphs[index].text.strip().endswith("PG4"):
            answer_index = index
            break

    if answer_index is None:
        return

    # Find a reference paragraph (ד. PG4) to copy its run formatting
    ref_paragraph = doc.paragraphs[answer_index]
    ref_runs = ref_paragraph.runs  # e.g. ('ד. ', size, None), ('PG', size, Courier New), ('4', size, Courier New)

    answers = [
        (answer_index + 1, "\u05d4", "PG5"),
        (answer_index + 2, "\u05d5", "PG6"),
    ]
    for paragraph_index, letter, pg_text in answers:
        paragraph = doc.paragraphs[paragraph_index]
        # Skip if already correctly formatted (runs match reference structure)
        if (
            len(paragraph.runs) >= 3
            and paragraph.runs[0].text == letter + ". "
            and paragraph.runs[0].font.size == ref_runs[0].font.size
            and paragraph.runs[1].text == "PG"
            and paragraph.runs[1].font.name == ref_runs[1].font.name
        ):
            continue
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        mark_rtl(paragraph)

        # Replicate the run structure of the reference paragraph
        if len(ref_runs) >= 3:
            # Run 1: label  (letter + ". ")
            r1 = paragraph.add_run(letter + ". ")
            r1.font.size = ref_runs[0].font.size
            r1.font.name = ref_runs[0].font.name
            rpr1 = r1._element.get_or_add_rPr()
            if rpr1.find(qn("w:rtl")) is None:
                rpr1.append(OxmlElement("w:rtl"))
            # Run 2: "PG"
            r2 = paragraph.add_run("PG")
            r2.font.size = ref_runs[1].font.size
            r2.font.name = ref_runs[1].font.name
            # Run 3: digit
            r3 = paragraph.add_run(pg_text[2])  # "5" or "6"
            r3.font.size = ref_runs[2].font.size
            r3.font.name = ref_runs[2].font.name
        else:
            # Fallback: single run
            run = paragraph.add_run(letter + ". " + pg_text)
            rpr = run._element.get_or_add_rPr()
            if rpr.find(qn("w:rtl")) is None:
                rpr.append(OxmlElement("w:rtl"))


QUESTION_ANCHORS = ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10."]


def add_page_breaks_before_questions(doc: Document) -> None:
    """Add w:pageBreakBefore to the first paragraph of every question except Q1."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for anchor in QUESTION_ANCHORS[1:]:  # skip "1."
            if text.startswith(anchor + " ") or text.startswith(anchor + "\t"):
                ppr = paragraph._p.get_or_add_pPr()
                pb = ppr.find(qn("w:pageBreakBefore"))
                if pb is None:
                    pb = OxmlElement("w:pageBreakBefore")
                    ppr.insert(0, pb)
                pb.set(qn("w:val"), "1")
                break


def replace_docx_images(rendered: list[tuple[str, Path, Path, object]]) -> Path:
    if not DOCX.exists():
        raise FileNotFoundError(f"Missing Word source document: {DOCX}")

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.before-image-update-{stamp}.docx"
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)
    used: set[int] = set()
    svg_by_rel_id = {}
    search_start = 0
    for anchor, svg_path, preview_png_path, width in rendered:
        option_question = None
        for question_number in (4, 6, 9, 10):
            if svg_path.stem.startswith(f"q{question_number:02d}_option_"):
                option_question = question_number
                break
        if option_question is not None:
            search_start = max(search_start, question_start_index(doc, option_question))

        slot_index, slot = find_inline_slot_after(doc, anchor, used, search_start)
        search_start = slot_index + 1
        clear_paragraph(slot)
        slot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = slot.add_run()
        run.add_picture(str(preview_png_path), width=width)
        blip = run._element.xpath(".//a:blip")[0]
        rel_id = blip.get(qn("r:embed"))
        svg_by_rel_id[rel_id] = svg_path

    # Do not normalize or rewrite Word text here.
    # The Word document is the single source of truth for question wording and
    # answer labels; this script should only refresh diagram payloads.

    output_docx = DOCX
    try:
        doc.save(output_docx)
    except PermissionError:
        output_docx = ROOT / "source.updated.docx"
        doc.save(output_docx)
        print(f"{DOCX} is locked; wrote {output_docx} instead")

    if EMBED_SVG_IN_DOCX:
        replace_package_images_with_svg(output_docx, svg_by_rel_id)
    cleanup_unused_media(output_docx)
    print(f"Updated {output_docx}")
    print(f"Backup {backup}")
    return output_docx


def cleanup_unused_media(docx_path: Path) -> None:
    tmp_docx = docx_path.with_suffix(".clean.tmp.docx")
    ns_rel = "http://schemas.openxmlformats.org/package/2006/relationships"
    ns_doc_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    ET.register_namespace("", ns_rel)

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        document_root = ET.fromstring(zin.read("word/document.xml"))
        used_rel_ids = {
            value
            for elem in document_root.iter()
            for attr, value in elem.attrib.items()
            if attr == f"{{{ns_doc_rel}}}embed"
        }

        rels_root = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
        kept_media = set()
        for rel in list(rels_root):
            target = rel.attrib.get("Target", "")
            rel_id = rel.attrib.get("Id")
            if target.startswith("media/"):
                if rel_id not in used_rel_ids:
                    rels_root.remove(rel)
                else:
                    kept_media.add("word/" + target)

        for item in zin.infolist():
            if item.filename.startswith("word/media/") and item.filename not in kept_media:
                continue
            if item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, ET.tostring(rels_root, encoding="utf-8", xml_declaration=True))
                continue
            zout.writestr(item, zin.read(item.filename))

    tmp_docx.replace(docx_path)


def replace_package_images_with_svg(docx_path: Path, svg_by_rel_id: dict[str, Path]) -> None:
    tmp_docx = docx_path.with_suffix(".tmp.docx")
    ns_rel = "http://schemas.openxmlformats.org/package/2006/relationships"
    ns_ct = "http://schemas.openxmlformats.org/package/2006/content-types"
    ns_doc_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    ET.register_namespace("", ns_rel)

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        document_root = ET.fromstring(zin.read("word/document.xml"))
        used_rel_ids = {
            value
            for elem in document_root.iter()
            for attr, value in elem.attrib.items()
            if attr == f"{{{ns_doc_rel}}}embed"
        }

        rels_xml = zin.read("word/_rels/document.xml.rels")
        rels_root = ET.fromstring(rels_xml)

        old_media_names = set()
        svg_payloads = {}
        for rel in list(rels_root):
            rel_id = rel.attrib.get("Id")
            old_target = rel.attrib["Target"]
            if old_target.startswith("media/") and rel_id not in used_rel_ids:
                old_media_names.add("word/" + old_target)
                rels_root.remove(rel)
                continue

            svg_path = svg_by_rel_id.get(rel_id)
            if svg_path is None and old_target in SHARED_MEDIA_IMAGES:
                svg_path = IMAGE_DIR / f"{SHARED_MEDIA_IMAGES[old_target]}.svg"
            if svg_path is None:
                continue
            old_media_names.add("word/" + old_target)
            new_target = str(Path(old_target).with_suffix(".svg")).replace("\\", "/")
            rel.attrib["Target"] = new_target
            svg_payloads["word/" + new_target] = svg_path.read_bytes()

        keep_media_names = {
            "word/" + rel.attrib["Target"]
            for rel in rels_root
            if rel.attrib.get("Target", "").startswith("media/")
        }

        content_types_xml = zin.read("[Content_Types].xml")
        content_types_root = ET.fromstring(content_types_xml)
        has_svg_default = any(
            child.tag == f"{{{ns_ct}}}Default" and child.attrib.get("Extension") == "svg"
            for child in content_types_root
        )
        if not has_svg_default:
            ET.SubElement(content_types_root, f"{{{ns_ct}}}Default", {
                "Extension": "svg",
                "ContentType": "image/svg+xml",
            })

        for item in zin.infolist():
            if (
                item.filename.startswith("word/media/")
                and item.filename not in keep_media_names
                and item.filename not in svg_payloads
            ):
                continue
            if item.filename in old_media_names:
                continue
            if item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, ET.tostring(rels_root, encoding="utf-8", xml_declaration=True))
                continue
            if item.filename == "[Content_Types].xml":
                zout.writestr(item, ET.tostring(content_types_root, encoding="utf-8", xml_declaration=True))
                continue
            zout.writestr(item, zin.read(item.filename))

        for filename, payload in svg_payloads.items():
            zout.writestr(filename, payload)

    tmp_docx.replace(docx_path)


def inspect_docx_image_slots(docx_path: Path) -> None:
    doc = Document(docx_path)
    rels = doc.part.rels
    for index, paragraph in enumerate(inline_image_paragraphs(doc), start=1):
        inline = paragraph._p.xpath(".//wp:inline")
        if not inline:
            print(f"{index:02}: non-inline image")
            continue
        extent = inline[0].find(qn("wp:extent"))
        blip = inline[0].xpath(".//a:blip")[0]
        rel_id = blip.get(qn("r:embed"))
        target = rels[rel_id].target_ref
        width = int(extent.get("cx")) / 914400
        height = int(extent.get("cy")) / 914400
        print(f"{index:02}: {target} {width:.2f}x{height:.2f} in")


def main() -> None:
    rendered = []
    for stem in dict.fromkeys(SHARED_MEDIA_IMAGES.values()):
        svg_path, _preview_png_path = render_tex(stem)
        print(f"Rendered {svg_path.relative_to(ROOT)}")

    for anchor, stem, width in QUESTION_IMAGES:
        svg_path, preview_png_path = render_tex(stem)
        rendered.append((anchor, svg_path, preview_png_path, width))
        print(f"Rendered {svg_path.relative_to(ROOT)}")

    updated_docx = replace_docx_images(rendered)
    inspect_docx_image_slots(updated_docx)


if __name__ == "__main__":
    main()
